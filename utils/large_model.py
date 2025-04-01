# utils/large_model.py

import os
import sys
import logging
from transformers import AutoTokenizer, AutoModelForSequenceClassification, pipeline
from huggingface_hub import HfApi
from docx import Document
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph  # 导入 Paragraph 类
from docx.table import _Cell  # 导入 Cell 类
from openpyxl import load_workbook
from openpyxl.styles import PatternFill
import requests

# 配置日志记录
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger('large_model')

MODEL_NAME = "alibaba-pai/pai-bert-base-zh-llm-risk-detection"  # 更新为目标模型
MODEL_PATH = os.path.join(os.getcwd(), "models")

# 全局变量用于缓存模型
_cached_classifier = None


def check_model_configured():
    """检查模型和分词器是否已下载并配置"""
    global MODEL_PATH  # 确保在函数开始时声明全局变量
    
    try:
        # 如果当前是打包后的环境，则尝试在打包目录中查找模型
        if getattr(sys, 'frozen', False):
            # 获取应用程序的根目录
            app_dir = os.path.dirname(sys.executable)
            bundled_model_dir = os.path.join(app_dir, 'models', MODEL_NAME.replace('/', '_'))
            logger.info(f"检查打包环境中的模型: {bundled_model_dir}")
            
            # 如果打包目录中存在模型，则使用打包目录中的模型
            if os.path.exists(bundled_model_dir):
                MODEL_PATH = os.path.join(app_dir, 'models')
                logger.info(f"使用打包目录中的模型: {MODEL_PATH}")
        
        model_dir = os.path.join(MODEL_PATH, MODEL_NAME.replace('/', '_'))
        logger.info(f"检查模型目录: {model_dir}")
        
        # 根据 alibaba-pai/pai-bert-base-zh-llm-risk-detection 模型调整所需文件列表
        required_files = ["pytorch_model.bin", "config.json", "tokenizer_config.json", "vocab.txt"]
        for file in required_files:
            file_path = os.path.join(model_dir, file)
            if not os.path.exists(file_path):
                logger.warning(f"缺少必要的模型文件: {file}")
                return False
        
        logger.info("模型已完整配置")
        return True
    except Exception as e:
        logger.error(f"检查模型配置时出错: {str(e)}")
        return False


def check_and_download_model(progress_callback):
    """
    检查并下载大模型。如果模型已配置完成，则直接返回。
    否则，下载模型和分词器，并通过 progress_callback 更新下载进度。
    """
    try:
        # 先检查模型是否已经配置好
        if check_model_configured():
            progress_callback("大模型已配置完成。")
            return
        
        api = HfApi()
        model_dir = os.path.join(MODEL_PATH, MODEL_NAME.replace('/', '_'))
        if not os.path.exists(model_dir):
            os.makedirs(model_dir)
        
        progress_callback("正在下载大模型，请稍候...")
        logger.info(f"开始下载模型到: {model_dir}")

        # 获取仓库中的所有文件
        try:
            repo_files = api.list_repo_files(repo_id=MODEL_NAME)
            total_files = len(repo_files)
            progress_callback(f"总共需要下载 {total_files} 个文件。")
            logger.info(f"需要下载的文件数量: {total_files}")
        except Exception as e:
            error_msg = f"获取模型文件列表失败: {str(e)}"
            logger.error(error_msg)
            progress_callback(error_msg)
            raise Exception(error_msg)

        # 假设使用 'main' 版本
        revision = "main"

        for idx, file_name in enumerate(repo_files, start=1):
            progress_callback(f"正在下载文件 {idx}/{total_files}: {file_name}")
            # 构建下载 URL
            download_url = f"https://huggingface.co/{MODEL_NAME}/resolve/{revision}/{file_name}"
            # 本地文件路径
            local_file_path = os.path.join(model_dir, file_name)
            temp_file_path = local_file_path + ".part"  # 临时文件名
            os.makedirs(os.path.dirname(local_file_path), exist_ok=True)
            
            # 下载文件
            try:
                with requests.get(download_url, stream=True) as r:
                    r.raise_for_status()
                    total_size = int(r.headers.get('content-length', 0))
                    downloaded_size = 0
                    chunk_size = 8192
                    # 为减少信号发送频率，设定一个阈值
                    update_threshold = chunk_size * 100  # 每下载100个块更新一次
                    next_update = update_threshold

                    with open(temp_file_path, 'wb') as f:
                        for chunk in r.iter_content(chunk_size=chunk_size):
                            if chunk:
                                f.write(chunk)
                                downloaded_size += len(chunk)
                                if downloaded_size >= next_update or downloaded_size == total_size:
                                    if total_size > 0:
                                        percent = int(downloaded_size * 100 / total_size)
                                    else:
                                        percent = 0
                                    progress_callback(f"下载进度: {percent}% ({downloaded_size}/{total_size} bytes)")
                                    next_update += update_threshold
                # 下载完成后重命名为目标文件名
                os.rename(temp_file_path, local_file_path)
                progress_callback(f"已下载文件 {idx}/{total_files}: {file_name}")
            except requests.exceptions.RequestException as e:
                error_msg = f"下载文件 {file_name} 时网络错误: {str(e)}"
                logger.error(error_msg)
                progress_callback(error_msg)
                if os.path.exists(temp_file_path):
                    os.remove(temp_file_path)
                raise Exception(error_msg)
            except Exception as e:
                error_msg = f"下载文件 {file_name} 时发生错误: {str(e)}"
                logger.error(error_msg)
                progress_callback(error_msg)
                if os.path.exists(temp_file_path):
                    os.remove(temp_file_path)
                raise Exception(error_msg)

        progress_callback("大模型下载完成并已配置。")
        logger.info("模型下载完成")
    except Exception as e:
        error_msg = f"下载大模型时发生错误：{str(e)}"
        logger.error(error_msg)
        raise Exception(error_msg)


def load_classifier(device='cpu'):
    """
    加载并缓存分类器，以避免重复加载模型。
    """
    global _cached_classifier
    
    try:
        # 确保模型已配置，这会正确设置MODEL_PATH
        if not check_model_configured():
            error_msg = "模型未配置完成，请先下载模型。"
            logger.error(error_msg)
            raise Exception(error_msg)
        
        if _cached_classifier is None:
            model_dir = os.path.join(MODEL_PATH, MODEL_NAME.replace('/', '_'))
            logger.info(f"开始加载模型: {model_dir}, 设备: {device}")
            
            try:
                tokenizer = AutoTokenizer.from_pretrained(model_dir)
                model = AutoModelForSequenceClassification.from_pretrained(model_dir)
                _cached_classifier = pipeline(
                    'text-classification',
                    model=model,
                    tokenizer=tokenizer,
                    device=0 if device == 'cuda' else -1,
                    max_length=128,
                    top_k=None,  # 使用 top_k=None 代替 return_all_scores=True
                    truncation=True  # 显式启用文本长度截断
                )
                logger.info("模型加载成功")
            except Exception as e:
                error_msg = f"加载模型失败: {str(e)}。请检查模型文件是否完整或重新下载模型。"
                logger.error(error_msg)
                raise Exception(error_msg)
        
        return _cached_classifier
    except Exception as e:
        error_msg = f"加载分类器时发生错误: {str(e)}"
        logger.error(error_msg)
        raise Exception(error_msg)


def analyze_files_with_model(file_paths, progress_callback, device='cpu', normal_threshold=0.8, other_threshold=0.1):
    """
    使用大模型分析多个文件内容，并根据风险等级进行标记。
    支持 .docx 和 .xlsx 文件。
    返回分析结果的列表，每个元素对应一个文件的统计信息。
    """
    try:
        # 检查模型是否已配置
        if not check_model_configured():
            error_msg = "大模型未配置，请先下载并配置大模型。"
            logger.error(error_msg)
            raise Exception(error_msg)

        # 加载模型
        progress_callback("加载模型中...")
        logger.info(f"开始分析文件，使用设备: {device}, 阈值: normal={normal_threshold}, other={other_threshold}")
        
        try:
            classifier = load_classifier(device)
        except Exception as e:
            error_msg = f"加载模型时发生错误：{str(e)}"
            logger.error(error_msg)
            progress_callback(error_msg)
            raise Exception(error_msg)

        results = []

        # 验证文件路径
        valid_file_paths = []
        for path in file_paths:
            if not os.path.exists(path):
                progress_callback(f"警告: 文件不存在 - {path}")
                logger.warning(f"文件不存在: {path}")
                continue
                
            file_type = os.path.splitext(path)[1].lower()
            if file_type not in ['.docx', '.xlsx']:
                progress_callback(f"警告: 不支持的文件类型 - {path}")
                logger.warning(f"不支持的文件类型: {path}")
                continue
                
            valid_file_paths.append(path)
            
        if not valid_file_paths:
            error_msg = "没有有效的可分析文件(.docx或.xlsx)"
            logger.error(error_msg)
            progress_callback(error_msg)
            raise Exception(error_msg)
            
        file_paths = valid_file_paths
        
        # 预先计算总处理项数
        total_items = 0
        items_per_file = []  # 存储每个文件的项数
        for file_path in file_paths:
            file_type = os.path.splitext(file_path)[1].lower()
            if file_type == '.docx':
                try:
                    doc = Document(file_path)
                    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
                    tables = [cell.text.strip() for table in doc.tables for row in table.rows for cell in row.cells if cell.text.strip()]
                    combined_texts = paragraphs + tables
                    if not combined_texts:
                        progress_callback(f"文件 {file_path} 中未找到任何可处理的段落或表格内容。")
                        logger.warning(f"文件 {os.path.basename(file_path)} 没有内容")
                        items = 0
                    else:
                        # 计算合并后的文本组数
                        merged_texts = merge_texts(combined_texts, min_length=60)
                        progress_callback(f"文件 {file_path} 共有 {len(paragraphs)} 个非空段落，{len(tables)} 个表格单元格，总合并后为 {len(merged_texts)} 个组。")
                        logger.info(f"文件 {os.path.basename(file_path)}: {len(merged_texts)} 个文本组")
                        items = len(merged_texts)
                except Exception as e:
                    progress_callback(f"处理 Word 文件 {file_path} 时发生错误：{str(e)}")
                    logger.error(f"处理Word文件错误: {os.path.basename(file_path)} - {str(e)}")
                    items = 0
            elif file_type == '.xlsx':
                try:
                    # 移除 read_only=True 以允许修改单元格样式
                    wb = load_workbook(file_path)
                    group_count = 0
                    for sheet in wb.worksheets:
                        columns = sheet.iter_cols(min_row=1, max_row=sheet.max_row, values_only=False)
                        for col in columns:
                            # 提取有值的单元格
                            cells = [cell for cell in col if cell.value]
                            if not cells:
                                continue
                            # 计算组数，每5个单元格为一组
                            group_count += (len(cells) + 4) // 5
                    progress_callback(f"文件 {file_path} 共有 {group_count} 个单元格组。")
                    logger.info(f"文件 {os.path.basename(file_path)}: {group_count} 个单元格组")
                    items = group_count
                except Exception as e:
                    progress_callback(f"处理 Excel 文件 {file_path} 时发生错误：{str(e)}")
                    logger.error(f"处理Excel文件错误: {os.path.basename(file_path)} - {str(e)}")
                    items = 0
            else:
                items = 0  # Unsupported file type
                progress_callback(f"文件 {file_path} 是不支持的类型，仅支持 .docx 和 .xlsx 文件。")
                logger.warning(f"不支持的文件类型: {file_path}")
            items_per_file.append(items)
            total_items += items

        if total_items == 0:
            error_msg = "未找到任何可处理的段落或单元格。"
            logger.error(error_msg)
            progress_callback(error_msg)
            raise Exception(error_msg)

        progress_callback(f"总共有 {total_items} 个待处理的段落或单元格组。")
        logger.info(f"总共有 {total_items} 个待处理项")

        # 初始化全局进度跟踪
        global_progress = {'processed': 0}

        for file_idx, file_path in enumerate(file_paths, start=1):
            result = {
                'file_path': file_path,
                'total_word_count': 0,
                'normal_count': 0,
                'low_vulgar_count': 0,
                'porn_count': 0,
                'other_risk_count': 0,
                'adult_count': 0,
                'new_file_path': ''
            }
            try:
                progress_callback(f"开始分析文件 {file_idx}/{len(file_paths)}: {os.path.basename(file_path)}")
                logger.info(f"开始分析文件: {os.path.basename(file_path)}")
                
                # 分析单个文件
                file_result = analyze_single_file_with_model(
                    file_path, classifier, progress_callback, global_progress, total_items,
                    normal_threshold, other_threshold)
                result.update(file_result)
                results.append(result)
                
                progress_callback(f"完成分析文件 {file_idx}/{len(file_paths)}: {os.path.basename(file_path)}\n")
                logger.info(f"完成分析文件: {os.path.basename(file_path)}")
            except Exception as e:
                error_msg = f"分析文件 {os.path.basename(file_path)} 时发生错误：{str(e)}"
                progress_callback(error_msg)
                logger.error(error_msg)
                results.append(result)  # 记录错误但继续

        progress_callback("所有文件分析完成！")
        logger.info("所有文件分析完成")
        return results
        
    except Exception as e:
        error_msg = f"分析文件过程中发生错误: {str(e)}"
        logger.error(error_msg)
        progress_callback(error_msg)
        raise Exception(error_msg)


def merge_texts(texts, min_length=60):
    """
    合并文本，直到合并后的文本长度达到最小长度。
    返回一个文本字符串列表。
    """
    merged = []
    current_group = ""
    for text in texts:
        if not text:
            continue
        if len(current_group) + len(text) + 1 <= min_length:
            if current_group:
                current_group += " " + text
            else:
                current_group = text
        else:
            if current_group:
                merged.append(current_group)
            current_group = text
    if current_group:
        merged.append(current_group)
    return merged


def analyze_single_file_with_model(file_path, classifier, progress_callback, global_progress, total_items, normal_threshold, other_threshold, batch_size=32):
    """
    使用大模型分析单个文件内容，并根据风险等级进行标记。
    支持 .docx 和 .xlsx 文件。
    返回分析结果的统计信息。
    """
    # 初始化统计信息
    total_word_count = 0
    normal_count = 0
    low_vulgar_count = 0
    porn_count = 0
    other_risk_count = 0
    adult_count = 0

    # 分析文件
    file_type = file_path.split('.')[-1].lower()
    if file_type == 'docx':
        try:
            doc = Document(file_path)
            paragraphs = [para for para in doc.paragraphs if para.text.strip()]
            tables = [cell for table in doc.tables for row in table.rows for cell in row.cells if cell.text.strip()]
            combined_texts = [para.text.strip() for para in paragraphs] + [cell.text.strip() for cell in tables]
            sources = paragraphs + tables

            if not combined_texts:
                progress_callback(f"文件 {file_path} 中未找到任何可处理的段落或表格内容。")
                return {
                    'total_word_count': total_word_count,
                    'normal_count': normal_count,
                    'low_vulgar_count': low_vulgar_count,
                    'porn_count': porn_count,
                    'other_risk_count': other_risk_count,
                    'adult_count': adult_count,
                    'new_file_path': ''
                }

            # 合并文本并保留源对象
            merged_texts_with_sources = merge_texts_with_sources(combined_texts, sources, min_length=60)
            progress_callback(f"文件 {file_path} 合并后共有 {len(merged_texts_with_sources)} 个文本组。")

            # 准备批量推理
            text_groups = [text_group for text_group, _ in merged_texts_with_sources]
            sources_groups = [source_list for _, source_list in merged_texts_with_sources]

            # 批量推理
            batch_labels_per_group = []
            for i in range(0, len(text_groups), batch_size):
                batch_texts = text_groups[i:i + batch_size]
                try:
                    batch_results = classifier(batch_texts)
                except Exception as e:
                    progress_callback(f"批量分析文本组时发生错误：{e}")
                    batch_results = [None] * len(batch_texts)  # 使用None表示分析失败

                for j, result in enumerate(batch_results):
                    if result is None:
                        group_label = "未知"
                    else:
                        # 根据每个文本的结果确定标签
                        group_label = determine_label_from_result(result, normal_threshold, other_threshold)
                    batch_labels_per_group.append(group_label)
                    # 更新全局进度
                    global_progress['processed'] += 1
                    percent = int((global_progress['processed'] / total_items) * 100)
                    progress_callback(f"进度: {percent}%")

            # 应用标签和统计
            label_color_mapping = {
                "低俗": RGBColor(0, 255, 0),      # 绿色
                "色情": RGBColor(255, 255, 0),    # 黄色
                "其他风险": RGBColor(255, 0, 0),  # 红色
                "成人": RGBColor(0, 0, 255)       # 蓝色
            }

            for group_idx, label in enumerate(batch_labels_per_group):
                if label in label_color_mapping:
                    color = label_color_mapping[label]
                    for source in sources_groups[group_idx]:
                        if isinstance(source, Paragraph):
                            # 仅修改包含问题的文本的运行
                            for run in source.runs:
                                if run.text.strip():
                                    run.font.color.rgb = color
                        elif isinstance(source, _Cell):
                            for paragraph in source.paragraphs:
                                for run in paragraph.runs:
                                    if run.text.strip():
                                        run.font.color.rgb = color
                    # 更新统计
                    if label == "低俗":
                        low_vulgar_count += 1
                    elif label == "色情":
                        porn_count += 1
                    elif label == "其他风险":
                        other_risk_count += 1
                    elif label == "成人":
                        adult_count += 1
                elif label == "正常":
                    normal_count += 1
                else:
                    pass  # 对于"未知"标签不做处理

                # 统计字数
                total_word_count += len(text_groups[group_idx])

            new_file_path = file_path.replace('.docx', '_analyzed.docx')
            doc.save(new_file_path)

            # 更新结果
            result = {
                'total_word_count': total_word_count,
                'normal_count': normal_count,
                'low_vulgar_count': low_vulgar_count,
                'porn_count': porn_count,
                'other_risk_count': other_risk_count,
                'adult_count': adult_count,
                'new_file_path': new_file_path
            }
            return result

        except Exception as e:
            progress_callback(f"处理 Word 文档时发生错误：{e}")
            raise

    elif file_type == 'xlsx':
        try:
            wb = load_workbook(file_path)
            for sheet in wb.worksheets:
                columns = sheet.iter_cols(min_row=1, max_row=sheet.max_row, values_only=False)
                for col in columns:
                    # 提取有值的单元格
                    cells = [cell for cell in col if cell.value]
                    if not cells:
                        continue

                    # 按每五个单元格分组
                    for i in range(0, len(cells), 5):
                        group_cells = cells[i:i + 5]
                        group_texts = [str(cell.value).strip() for cell in group_cells]
                        concatenated_text = ''.join(group_texts)
                        total_word_count += len(concatenated_text)

                        # 分段处理，如果超过128个字符
                        segments = []
                        for j in range(0, len(concatenated_text), 128):
                            segment = concatenated_text[j:j + 128]
                            segments.append(segment)

                        # 如果没有分段（即总长度 <= 128），仍然需要进行分析
                        if not segments:
                            segments = [concatenated_text]

                        # 批量推理
                        try:
                            batch_results = classifier(segments)
                        except Exception as e:
                            progress_callback(f"批量分析单元格组时发生错误：{e}")
                            batch_results = [None] * len(segments)  # 使用None表示分析失败

                        group_labels = []
                        for result in batch_results:
                            if result is None:
                                group_labels.append("未知")
                            else:
                                group_label = determine_label_from_result(result, normal_threshold, other_threshold)
                                group_labels.append(group_label)
                            # 更新全局进度
                            global_progress['processed'] += 1
                            percent = int((global_progress['processed'] / total_items) * 100)
                            progress_callback(f"进度: {percent}%")

                        # 判断是否需要标记
                        mark_group = False
                        for label in group_labels:
                            if label != "正常" and label != "未知":
                                mark_group = True
                                break

                        if mark_group:
                            # 根据第一个非正常标签进行标记
                            group_label = next((label for label in group_labels if label != "正常" and label != "未知"), "其他风险")

                            # 应用标签和标记
                            fill_color_mapping = {
                                "低俗": '00FF00',      # 绿色
                                "色情": 'FFFF00',      # 黄色
                                "其他风险": 'FF0000',  # 红色
                                "成人": '0000FF'       # 蓝色
                            }

                            fill_color = fill_color_mapping.get(group_label, None)

                            if fill_color:
                                fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type='solid')
                                for cell in group_cells:
                                    cell.fill = fill
                                    # 更新统计
                                    if group_label == "低俗":
                                        low_vulgar_count += 1
                                    elif group_label == "色情":
                                        porn_count += 1
                                    elif group_label == "其他风险":
                                        other_risk_count += 1
                                    elif group_label == "成人":
                                        adult_count += 1
                        else:
                            normal_count += len(group_cells)

            new_file_path = file_path.replace('.xlsx', '_analyzed.xlsx')
            wb.save(new_file_path)

            # 更新结果
            result = {
                'total_word_count': total_word_count,
                'normal_count': normal_count,
                'low_vulgar_count': low_vulgar_count,
                'porn_count': porn_count,
                'other_risk_count': other_risk_count,
                'adult_count': adult_count,
                'new_file_path': new_file_path
            }
            return result

        except Exception as e:
            progress_callback(f"处理 Excel 文档时发生错误：{e}")
            raise

    else:
        raise ValueError("仅支持 .docx 和 .xlsx 文件")

    # 构建结果统计信息
    result = {
        'total_word_count': total_word_count,
        'normal_count': normal_count,
        'low_vulgar_count': low_vulgar_count,
        'porn_count': porn_count,
        'other_risk_count': other_risk_count,
        'adult_count': adult_count,
        'new_file_path': new_file_path
    }

    return result


def determine_label_from_result(result, normal_threshold, other_threshold):
    """
    根据模型输出结果和阈值，确定文本组的标签。
    """
    # 查找"正常"类别的分数
    normal_score = next((item['score'] for item in result if item['label'] == "正常"), 0)
    if normal_score >= normal_threshold:
        return "正常"
    else:
        # 获取其他类别中分数最高的类别
        other_scores = [item for item in result if item['label'] != "正常"]
        if not other_scores:
            return "其他风险"
        max_other = max(other_scores, key=lambda x: x['score'])
        if max_other['score'] >= other_threshold:
            return max_other['label']
        else:
            return "其他风险"


def merge_texts_with_sources(texts, sources, min_length=60):
    """
    合并文本，同时记录每个文本组对应的源（段落或表格单元格）。
    返回一个列表的元组： (text_group, list_of_sources)
    """
    merged = []
    current_group = ""
    current_sources = []
    for text, source in zip(texts, sources):
        if len(current_group) + len(text) + 1 <= min_length:
            if current_group:
                current_group += " " + text
            else:
                current_group = text
            current_sources.append(source)
        else:
            if current_group:
                merged.append((current_group, current_sources.copy()))
            current_group = text
            current_sources = [source]
    if current_group:
        merged.append((current_group, current_sources.copy()))
    return merged
