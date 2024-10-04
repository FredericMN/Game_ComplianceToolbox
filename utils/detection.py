import os
from docx import Document
from docx.shared import RGBColor  # 导入 RGBColor
from openpyxl import load_workbook
from openpyxl.styles import PatternFill

# 获取当前工作目录
CURRENT_DIR = os.getcwd()  # 使用当前工作目录

VIOLENT_WORDS_FILE = 'violent_words.txt'
INDUCING_WORDS_FILE = 'inducing_words.txt'

VIOLENT_WORDS_PATH = os.path.join(CURRENT_DIR, VIOLENT_WORDS_FILE)
INDUCING_WORDS_PATH = os.path.join(CURRENT_DIR, INDUCING_WORDS_FILE)

# 获取当前文件所在目录
CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))

# 读取词汇列表
def load_words(file_name, default_words):
    if os.path.exists(file_name):
        with open(file_name, 'r', encoding='utf-8') as f:
            content = f.read()
            # 处理中英文逗号，并分割词汇
            words = content.replace('，', ',').split(',')
            words = [word.strip() for word in words if word.strip()]
            return words
    else:
        # 文件不存在时，使用默认词汇并创建文件
        save_words(file_name, default_words)
        return default_words

# 保存词汇列表
def save_words(file_name, words):
    with open(file_name, 'w', encoding='utf-8') as f:
        f.write(','.join(words))

# 初始化词汇列表
def initialize_words():
    violent_words = load_words(VIOLENT_WORDS_PATH, [
        '杀', '死', '血', '屠', '戮', '亡', '残暴', '屠杀', '屠宰', '暴力',
        '虐待', '强奸', '抢劫', '恐怖', '爆炸', '枪击', '刺杀', '劫持', '恐吓', '袭击'
    ])
    inducing_words = load_words(INDUCING_WORDS_PATH, [
        '吸粉', '毒', '吸烟', '抽烟', '酒', '烟草', '犯罪', '妈的', '诱导', '诈骗', '盗取',
        '赌博', '色情', '偷取', '赌博', '毒品', '违法', '非法', '诱骗', '欺诈', '煽动'
    ])
    return violent_words, inducing_words

# 检测文档的函数
def detect_language(file_path, violent_words, inducing_words):
    violent_count = 0
    inducing_count = 0
    total_word_count = 0
    file_type = file_path.split('.')[-1].lower()

    if file_type == 'docx':
        doc = Document(file_path)
        for para in doc.paragraphs:
            text = para.text
            total_word_count += len(text.replace(" ", "").replace("\n", ""))  # 统计总字数（去掉空格和换行符）

            new_runs = []
            current_run_text = ""

            for char in text:
                highlighted = False

                # 检测是否为暴力词汇
                if char in violent_words:
                    if current_run_text:
                        new_runs.append({'text': current_run_text, 'color': None})
                        current_run_text = ""
                    new_runs.append({'text': char, 'color': 'FF0000'})  # 红色字体
                    violent_count += 1
                    highlighted = True

                # 检测是否为诱导词汇
                elif char in inducing_words:
                    if current_run_text:
                        new_runs.append({'text': current_run_text, 'color': None})
                        current_run_text = ""
                    new_runs.append({'text': char, 'color': '00FF00'})  # 绿色字体
                    inducing_count += 1
                    highlighted = True

                if not highlighted:
                    current_run_text += char

            if current_run_text:
                new_runs.append({'text': current_run_text, 'color': None})

            # 清除原有文本，逐字添加新 Run
            para.clear()
            for item in new_runs:
                run = para.add_run(item['text'])
                if item['color']:
                    # 设置字体颜色
                    run.font.color.rgb = RGBColor(int(item['color'][0:2], 16),
                                                 int(item['color'][2:4], 16),
                                                 int(item['color'][4:6], 16))

        new_file_path = file_path.replace('.docx', '_modified.docx')
        doc.save(new_file_path)

    elif file_type == 'xlsx':
        wb = load_workbook(file_path)
        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        cell_text = str(cell.value)
                        total_word_count += len(cell_text.replace(" ", "").replace("\n", ""))  # 统计总字数

                        if any(word in cell_text for word in violent_words):
                            cell.fill = PatternFill(start_color='FF4C4C', end_color='FF4C4C', fill_type='solid')  # 红色
                            violent_count += sum(cell_text.count(word) for word in violent_words)

                        if any(word in cell_text for word in inducing_words):
                            cell.fill = PatternFill(start_color='4CFF4C', end_color='4CFF4C', fill_type='solid')  # 绿色
                            inducing_count += sum(cell_text.count(word) for word in inducing_words)

        new_file_path = file_path.replace('.xlsx', '_modified.xlsx')
        wb.save(new_file_path)

    else:
        raise ValueError("Unsupported file type. Please select a .docx or .xlsx file.")

    return violent_count, inducing_count, total_word_count, new_file_path

# 更新词汇列表函数
def update_words(violent_input, inducing_input):
    # 处理中英文逗号，并分割
    violent_input = violent_input.replace('，', ',')
    inducing_input = inducing_input.replace('，', ',')

    # 更新词汇列表
    violent_words = [word.strip() for word in violent_input.split(',') if word.strip()]
    inducing_words = [word.strip() for word in inducing_input.split(',') if word.strip()]

    # 保存到文件
    save_words(VIOLENT_WORDS_PATH, violent_words)
    save_words(INDUCING_WORDS_PATH, inducing_words)

    return violent_words, inducing_words
